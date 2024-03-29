import pandas as   pd
import time
import datetime
import os
from mysqlmodel import mysql_model
from  docx import Document
import re
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import decimal  


re1="[Y|Z](\d{4})-\d{5}$"
re2="[Y|Z]\d{4}-(\d{5})$"
class yugum():
    def __init__(self):
        pass
    def do(self,pathx,pathw,dataname,us,pwd,host1,ch,mu,tablename,searchlist=None):
        model(pathx,pathw,dataname,us,pwd,host1,ch,mu,tablename,searchlist)

def model(pathx,pathw,dataname,us,pwd,host1,ch,mu,tablename,searchlist=None):
    readexcel_r=readexcel(pathx)
    ls2=['fd42_hx','fd42_lx','fd42_lc','fd42_td','fd42_y','fd42_jg']
    ls3=['fd_dx1','fd_dx2','fd_dx3','fd_dx4','fd_dx5']
    df=getlistl(dataname,us,pwd,host1,ch,mu,tablename,readexcel_r,searchlist)
    df['fd37']=df['fd37'].apply(lambda x : x if str(x).find("股份有限公司")>-1 or str(x)=='上海银行' or str(x)=="南京银行" else str(x)+"股份有限公司")
    df['fd38']=df['fd38'].apply(lambda x:str(x)+"支行" if str(x).find("支行")<0 and  str(x).find("分行")<0 else str(x) )
    df['fd1_year']=df['fd1'].apply(lambda x: re.findall(re1,x)[0])
    df['fd1_2']=df['fd1'].apply(lambda x: re.findall(re2,x)[0])
    df['fd10_1']=df['fd10'].apply(lambda x :riq2(x))
    df['fd10_2']=df['fd10'].apply(lambda x:qzhuan(x))
    df['fd20_z']=df['fd20'].apply(lambda x :zhuandaxie(int(x*10000)))
    #df['fdzq']=df['fdzq'].apply(lambda x : str(x).split("，")[0])
    df['landfeature']=df['landfeature'].apply(lambda x : '国有' if x =='' else x)
    df['fd16']=df['fd16'].apply(lambda x : '出让' if x =='' else x)
    df['fd_dz_1']=df['fd3'].apply(lambda x :fen1(x))
    df['fd_dz_2']=df['fd3'].apply(lambda x :fen2(x))
    df1=df[['fd38','fd1_year','fd1_2','fd10_1','fd10_2','fd20_z','fdzq','buildingname','landfeature','fd37','district','fd3','fd16','fd18','fd21','fd4','fd1','fd20','fd15','credentialsno','fd_dz_2','fd_dz_1','fdjing']]
    df1['title']=df1['fd1'].apply(lambda  x: '预评估报告' if x[0]=='Y' else '询价报告')
    #df1['fh']=df1.apply(relfh,axis=1)
    df1['fh']=df1['fd37'].apply(lambda x : relfh(x))

    r1=df1.to_dict("records")
    
    print(r1)
    #print(df)
    del df
    del df1

    for rr in r1:
        if rr['fdzq']:
            count=rr['fdzq'].count('-')
            if count <5:
                dc=5-count
                for idc in range(dc):
                    if rr['fdzq']:
                        rr['fdzq']=rr['fdzq']+r"-/"
                    else:
                        rr['fdzq']=r'/-/'
            ls_l=rr['fdzq'].split('-')
            for ii,ls_l2 in enumerate(ls_l):
                rr[ls2[ii]]=ls_l2
        rr.pop('fdzq')
        if not rr['credentialsno']=="":
            count=rr['credentialsno'].count('-')
            if count <4:
                dc=4-count
                for idc in range(dc):
                    if rr['credentialsno']:
                        rr['credentialsno']=rr['credentialsno']+r"-/"
                    else:
                        rr['credentialsno']=r'/-/'
            ls_l=rr['credentialsno'].split('-')
            for ii,ls_l2 in enumerate(ls_l):
                rr[ls3[ii]]=ls_l2
            if rr['fd_dx5']==r"/":
                rr['fd_dx5']=0
        rr.pop('credentialsno')
        if not rr.__contains__("fd_dx1"):
            if rr['fd37']=="南京银行":
                print(rr["fdjing"])
                if rr["fdjing"]:

                    rr["fdjing_1"]=zhuandaxie(int(rr["fdjing"])*10000)
                    rr["fdjing_2"]=rr["fdjing"]
                    print(rr['fdjing_1'])
                pathw1=os.path.join(pathw,r"yugu\mu\模板3.docx")
            else:
                pathw1=os.path.join(pathw,r"yugu\mu\模板1.docx")
             #r"yugu\mu\模板1.docx"
        else:
            if not rr["fd_dx4"] ==r"/":
                if re.match(r"^\d+(\.\d{1,5})*$",rr['fd_dx4']):
                    rr["fd20_1_2"]=int(rr["fd20"])+int(rr['fd_dx4'])
                    rr["fd20_1_1"]=zhuandaxie(rr["fd20_1_2"]*10000)
                    rr["fd_dx4"]=format(float(rr["fd_dx4"]),'.2f')
                    rr["fd20_1_2"]=format(rr["fd20_1_2"],'.2f')
                    rr['fd3']="及".join([rr['fd3'],rr['fd_dx1']])
            pathw1=os.path.join(pathw,r"yugu\mu\模板2.docx")
        tiword(pathw1,rr,rr['fd1'])
        

                
def relfh(x):
  
    if x.find("工商")>-1:
        return "上海市"
    elif x=='上海银行':
        return ""
    else:
        return "上海分行"

def fen1(fd3):
    a1=fd3.find("弄")
    a2=fd3.find("号")
    a3=fd3.find("苑")
    a4=fd3.find("园")
    if a1>-1:
        return fd3[:a1+1]
    elif a2>-1:
        if a2>a3 and a2 >a4 :
            return fd3[:a2+1]
        else:
            if a3>a4:
                return fd3[:a3+1]
            else:
                return fd3[:a4+1] if a4>-1 else ""

def fen2(fd3):

    a1=fd3.find("弄")
    a2=fd3.find("号")
    a3=fd3.find("苑")
    a4=fd3.find("园")
    if a1>-1:
        return fd3[a1+1:]
    elif a2>-1:
        if a2>a3 and a2 >a4 :
            return fd3[a2+1:]
        else:
            if a3>a4:
                return fd3[a3+1:]
            else:
                return fd3[a4+1:] if a4>-1 else ""

def riq2(riq):
    r=""
    lp=[]
    if riq:
        
        rdatetime=datetime.datetime.strptime(str(riq),"%Y-%m-%d")
        lp.append(str(rdatetime.year))
        lp.append("年")
        lp.append(str(rdatetime.month))
        lp.append("月")
        lp.append(str(rdatetime.day))
        lp.append("日")
        r="".join(lp)
    return r

def tiword(path,dic,flname):
    flpathsave=path[:path.rfind("\\")+1]+dic['title']+flname+".docx"
    if os.path.exists(flpathsave):
        #path=flpathsave
        os.remove(flpathsave)
    print(path)
    docxx=Document(path)
    for para in docxx.paragraphs:
        for i in range(len(para.runs)):
            print(para.runs[i].text)
            if  dic["fd37"]=="南京银行":
                if para.runs[i].text.find("fd20]")>-1 :

                    para.runs[i].text=para.runs[i].text.replace("fd20]",str(dic["fd20"]))
                if para.runs[i].text.find("fdjing_1]")>-1:

                    para.runs[i].text=para.runs[i].text.replace("fdjing_1]",str(dic["fdjing_1"]))
            # if para.runs[i].text.find('三、价值时点：')>-1:
            #     para.runs[i].text='三、价值时点：'+dic['fd10_1']
            # if para.runs[i].text.find('title')>-1:
            #     para.runs[i].text=para.runs[i].text.replace("title",dic['title'])
            for k,v in dic.items():
                para.runs[i].text=para.runs[i].text.replace("["+str(k)+"]",str(v))
    table=docxx.tables
    for otable in table:
        rows_num=len(otable.rows)
        columns_num=len(otable.columns)
        for j in range(rows_num):
            for ll in range(columns_num):

                for k,v in dic.items():
                    sp=otable.cell(j,ll).text
                    otable.cell(j,ll).text=sp.replace("["+str(k)+"]",str(v))
                run = otable.cell(j,ll).paragraphs[0].runs
                for rrr in run :
                    if dic["fd37"]=="南京银行":
                        rrr.font.name = u'宋体'
                        r = rrr._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    else:
                        rrr.font.name = u'仿宋_gb2312'
                        r = rrr._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_gb2312')
                otable.cell(j,ll).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    docxx.save(flpathsave)


def getlistl(dataname,us,pwd,host1,ch,mu,tablename,list_l,searchlist):
    mysqlm=mysql_model(dataname,us,pwd,host1,ch)
    
    mresult=mysqlm.selectmul(mu,tablename,list_l,searchlist)
    
    df=pd.DataFrame(mresult,columns=searchlist)
    return df


def readexcel(path):
    df=pd.read_excel(path,header=0,sheet_name = 0,usecols='A')
    
    df.fillna("",inplace=True)
    pr=df.values.tolist()
    
    pr1=[x[0] for x in pr]
  
    return [pr1]



def  qzhuan(riq):
    r=[]
    riq1=str(riq)
    if riq :
        f=riq1.split("-")
        lp=["〇","一","二","三","四","五","六","七","八","九"]
        if len(f)==3:
            for i in f[0]:
                r.append(lp[int(i)])
            r.append("年")
            if len(f[1])>1 and f[1][0] !="0":
                if f[1][0] !="1":
                    r.append(lp[int(f[1][0])])
                r.append("十")
            if int(f[1][-1:]) !=0:
                r.append(lp[int(f[1][-1:])])
            r.append("月")
            if len(f[2])>1 and f[2][0] !="0":
                if f[2][0] !="1":
                    r.append(lp[int(f[2][0])])
                r.append("十")

            if int(f[2][-1:]) !=0:
                r.append(lp[int(f[2][-1:])])
            r.append("日")
    return ''.join(r)


def zhuandaxie(nums):
    nums =str(nums)
    da1=["零","壹","贰","叁","肆","伍","陆","柒","捌","玖"]
    da2=["拾","佰","仟","万","亿","兆"]
    dax=[]
    p=0
    if  is_number(nums):
        l=len(nums)  #长度
        
        if l<20:
            for i in range(0,l):
                ii=l-1-i
                nm=int(nums[i:i+1])
                if i<l-1 and nm==0:
                    if int(nums[i+1:i+2]) >0:
                        if ii %4 ==0:
                            if i-1>-1:
                                if int(nums[i-1:i]) >0:
                                    pass
                                else:
                                    dax.append(da1[0])
                        else:

                            dax.append(da1[0])
                elif nm==0 :
                    pass
                else:
                    dax.append(da1[nm])
                #if (ii+4) % 4 ==0  and i>0 :
                if ii==4  :
                    if nm>0:
                        dax.append(da2[3]) 
                    elif l<12:
                        for p in range(0,l-ii-1):
                            if int(nums[i-p-1:i])>0:
                                dax.append(da2[3]) 
                                break
                elif ii==12 :
                    dax.append(da2[3]) 
                elif ii==8:
                    dax.append(da2[4])
                elif ii== 16:
                    dax.append(da2[5])
                elif nm >0 and ii >0:
                    ll=(ii+4)%4
                    dax.append(da2[ll-1])
        return "".join(dax)
    else:
        return "数字太长"

def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        pass
 
    # try:
    #     import unicodedata
    #     unicodedata.numeric(s)
    #     return True
    # except (TypeError, ValueError):
    #     pass
 
    return False
if __name__ == "__main__":
    pp=os.path.join(os.path.abspath('.'),r"yugu\xlsx\预估报告编号.xlsx")
    p2=os.path.abspath('.')
    data_n="im2006"
    us="user"#"user"
    pwd="7940"#"7940"
    host="192.168.1.3"#"192.168.1.3"
    ch="utf8"
    mu=["fd1"]
    tablename="reports"
    searchlist=["district","fd3","fd1","buildingname","fdzq","fd4","fd18","fd15","fd20","fd21","fd10","fd37","fd38","fd16", "landfeature","credentialsno","fdjing"]
    model(pp,p2,data_n,us,pwd,host,ch,mu,tablename,searchlist)
    

    #print("".join(zhuandaxie(str(101500))))
